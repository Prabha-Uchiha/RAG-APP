import streamlit as st
import pandas as pd
import numpy as np
from typing import List, Optional, Tuple
import tempfile
import os
from pathlib import Path
import torch

# Core RAG imports - UPDATED IMPORTS
from langchain_core.documents import Document as LangchainDocument # New Core Document
from langchain_text_splitters import RecursiveCharacterTextSplitter # New splitter package
from langchain_community.vectorstores import FAISS # FAISS remains in community
from langchain_community.vectorstores.utils import DistanceStrategy
# Using the dedicated HuggingFace integration package
from langchain_huggingface import HuggingFaceEmbeddings
from transformers import AutoTokenizer, AutoModelForCausalLM, BitsAndBytesConfig, pipeline

# Document processing imports
import fitz # PyMuPDF for PDF processing
import docx # python-docx for Word documents

# Configuration
EMBEDDING_MODEL_NAME = "thenlper/gte-small"
READER_MODEL_NAME = "HuggingFaceH4/zephyr-7b-beta"

# Markdown separators for chunking
MARKDOWN_SEPARATORS = [
    "\n#{1,6} ",
    "```\n",
    "\n***+\n",
    "\n---+\n",
    "\n___+\n",
    "\n\n",
    "\n",
    " ",
    "",
]

class RAGSystem:
    def __init__(self):
        self.embedding_model = None
        self.vector_database = None
        self.reader_llm = None
        self.tokenizer = None
        self.prompt_template = None
        self.documents = []
        
    @st.cache_resource
    def initialize_models(_self):
        """Initialize embedding model and reader LLM"""
        try:
            # Initialize embedding model - Updated to use langchain_huggingface import
            _self.embedding_model = HuggingFaceEmbeddings(
                model_name=EMBEDDING_MODEL_NAME,
                model_kwargs={"device": "cpu"},  # Use CPU for free deployment
                encode_kwargs={"normalize_embeddings": True}
            )
            
            # Initialize tokenizer for reader model
            _self.tokenizer = AutoTokenizer.from_pretrained(READER_MODEL_NAME)
            if _self.tokenizer.pad_token is None:
                _self.tokenizer.pad_token = _self.tokenizer.eos_token
            
            # Configure quantization for Zephyr model (for memory efficiency)
            # This loads the model in 4-bit (NF4) to save VRAM.
            bnb_config = BitsAndBytesConfig(
                load_in_4bit=True,
                bnb_4bit_use_double_quant=True,
                bnb_4bit_quant_type="nf4",
                bnb_4bit_compute_dtype=torch.bfloat16,
            )
            
            # Load the Zephyr model with quantization
            model = AutoModelForCausalLM.from_pretrained(
                READER_MODEL_NAME, 
                quantization_config=bnb_config,
                device_map="auto",
                trust_remote_code=True
            )
            
            _self.reader_llm = pipeline(
                "text-generation",
                model=model,
                tokenizer=_self.tokenizer,
                do_sample=True,
                temperature=0.2,
                repetition_penalty=1.1,
                return_full_text=False,
                max_new_tokens=500,
            )
            
            # Create prompt template for Zephyr model
            system_message = """Using the information contained in the context, give a comprehensive answer to the question.
Respond only to the question asked, response should be concise and relevant to the question.
Provide the number of the source document when relevant.
If the answer cannot be deduced from the context, do not give an answer."""

            prompt_in_chat_format = [
                {"role": "system", "content": system_message},
                {"role": "user", "content": """Context:
{context}
---
Now here is the question you need to answer.
Question: {question}"""},
            ]
            
            _self.prompt_template = _self.tokenizer.apply_chat_template(
                prompt_in_chat_format, 
                tokenize=False, 
                add_generation_prompt=True
            )
            
            return True
        except Exception as e:
            st.error(f"Error initializing models: {str(e)}")
            return False
    
    def extract_text_from_file(self, uploaded_file) -> str:
        """Extract text from uploaded files"""
        try:
            if uploaded_file.type == "text/plain":
                return str(uploaded_file.read(), "utf-8")
            elif uploaded_file.type == "application/pdf":
                # Save uploaded file temporarily
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_file_path = tmp_file.name
                
                # Extract text from PDF
                doc = fitz.open(tmp_file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                os.unlink(tmp_file_path)
                return text
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # Handle Word documents
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_file_path = tmp_file.name
                
                doc = docx.Document(tmp_file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                os.unlink(tmp_file_path)
                return text
            else:
                st.error(f"Unsupported file type: {uploaded_file.type}")
                return ""
        except Exception as e:
            st.error(f"Error extracting text: {str(e)}")
            return ""
    
    def split_documents(self, chunk_size: int, documents: List[LangchainDocument]) -> List[LangchainDocument]:
        """Split documents into chunks - Updated to use langchain_text_splitters"""
        try:
            # Use the tokenizer from the embedding model for accurate chunking
            tokenizer = AutoTokenizer.from_pretrained(EMBEDDING_MODEL_NAME) 
            text_splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(
                tokenizer,
                chunk_size=chunk_size,
                chunk_overlap=int(chunk_size / 10),
                add_start_index=True,
                strip_whitespace=True,
                separators=MARKDOWN_SEPARATORS,
            )
            
            docs_processed = []
            for doc in documents:
                docs_processed += text_splitter.split_documents([doc])
            
            # Remove duplicates
            unique_texts = {}
            docs_processed_unique = []
            for doc in docs_processed:
                if doc.page_content not in unique_texts:
                    unique_texts[doc.page_content] = True
                    docs_processed_unique.append(doc)
            
            return docs_processed_unique
        except Exception as e:
            st.error(f"Error splitting documents: {str(e)}")
            return []
    
    def create_vector_database(self, docs_processed: List[LangchainDocument]):
        """Create FAISS vector database from processed documents"""
        try:
            if not self.embedding_model:
                st.error("Embedding model not initialized")
                return False
            
            self.vector_database = FAISS.from_documents(
                docs_processed, 
                self.embedding_model, 
                distance_strategy=DistanceStrategy.COSINE
            )
            return True
        except Exception as e:
            st.error(f"Error creating vector database: {str(e)}")
            return False
    
    def answer_with_rag(self, question: str, num_retrieved_docs: int = 5) -> Tuple[str, List[str]]:
        """Generate answer using RAG pipeline"""
        try:
            if not self.vector_database:
                return "No knowledge base available. Please upload documents first.", []
            
            # Retrieve relevant documents
            relevant_docs = self.vector_database.similarity_search(
                query=question, 
                k=num_retrieved_docs
            )
            
            # Extract document contents and metadata
            relevant_docs_text = []
            context_parts = []
            for i, doc in enumerate(relevant_docs):
                relevant_docs_text.append(doc.page_content)
                context_parts.append(f"Document {i} (Source: {doc.metadata.get('source', 'Unknown')}):::\n{doc.page_content}\n\n")
            
            # Build context
            context = "\nExtracted documents:\n" + "".join(context_parts)
            
            # Create final prompt using the template
            # We replace the placeholders in the prompt_template
            final_prompt = self.prompt_template.format(context=context, question=question)

            # Generate answer
            if self.reader_llm:
                # We use the generated final_prompt as the input to the pipeline
                response = self.reader_llm(final_prompt)
                # The response structure is often a list of dictionaries with 'generated_text'
                answer = response[0]["generated_text"].strip()
            else:
                answer = "Reader model not available. Showing retrieved context only."
            
            return answer, relevant_docs_text
        except Exception as e:
            # This catches issues like max context length exceeded for the prompt
            return f"Error generating answer: {str(e)}", []

def main():
    st.set_page_config(
        page_title="Advanced RAG System",
        page_icon="",
        layout="wide"
    )
    
    st.title("Advanced RAG System")
    st.markdown("Upload documents, build a knowledge base, and ask questions!")
    
    # Initialize RAG system
    if 'rag_system' not in st.session_state:
        st.session_state.rag_system = RAGSystem()
    
    rag_system = st.session_state.rag_system
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("Configuration")
        
        # Model initialization
        st.subheader("Model Setup")
        if st.button("Initialize Models", type="primary"):
            with st.spinner("Loading models..."):
                success = rag_system.initialize_models()
                if success:
                    st.success("Models loaded successfully!")
                else:
                    st.error("Failed to load models")
        
        # Chunking parameters
        st.subheader("Chunking Parameters")
        chunk_size = st.slider("Chunk Size (tokens)", 100, 1000, 512)
        
        # Retrieval parameters
        st.subheader("Retrieval Parameters")
        num_docs = st.slider("Number of documents to retrieve", 1, 10, 5)
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("Document Upload & Processing")
        
        # File upload
        uploaded_files = st.file_uploader(
            "Upload documents",
            accept_multiple_files=True,
            type=['txt', 'pdf', 'docx'],
            help="Upload PDF, DOCX, or TXT files to build your knowledge base"
        )
        
        if uploaded_files:
            st.subheader("Uploaded Files")
            for file in uploaded_files:
                st.write(f"📄 {file.name} ({file.type})")
        
        # Process documents button
        if st.button("Process Documents", type="primary") and uploaded_files:
            with st.spinner("Processing documents..."):
                # Extract text from files
                raw_documents = []
                for file in uploaded_files:
                    # Note: Using LangchainDocument (now from langchain_core)
                    text = rag_system.extract_text_from_file(file)
                    if text:
                        raw_documents.append(
                            LangchainDocument(
                                page_content=text,
                                metadata={"source": file.name}
                            )
                        )
                
                if raw_documents:
                    # Split documents into chunks
                    docs_processed = rag_system.split_documents(chunk_size, raw_documents)
                    st.session_state.docs_processed = docs_processed
                    
                    # Create vector database
                    if rag_system.embedding_model:
                        success = rag_system.create_vector_database(docs_processed)
                        if success:
                            st.success(f"Successfully processed {len(docs_processed)} chunks from {len(uploaded_files)} documents!")
                            st.session_state.knowledge_base_ready = True
                            
                            # Show statistics
                            st.subheader("Processing Statistics")
                            st.metric("Total Chunks", len(docs_processed))
                            st.metric("Average Chunk Length", 
                                    int(np.mean([len(doc.page_content) for doc in docs_processed])))
                        else:
                            st.error("Failed to create vector database")
                    else:
                        st.warning("Please initialize models first")
                else:
                    st.error("No valid documents found")
        
        # Show processed chunks
        if 'docs_processed' in st.session_state:
            st.subheader("Document Chunks Preview")
            if st.session_state.docs_processed:
                chunk_idx = st.selectbox(
                    "Select chunk to preview", 
                    range(len(st.session_state.docs_processed)),
                    format_func=lambda x: f"Chunk {x+1}"
                )
                
                selected_chunk = st.session_state.docs_processed[chunk_idx]
                st.text_area(
                    f"Chunk {chunk_idx + 1} Content",
                    selected_chunk.page_content,
                    height=200,
                    disabled=True
                )
                st.json(selected_chunk.metadata)
    
    with col2:
        st.header("Query Interface")
        
        # Check if knowledge base is ready
        knowledge_base_ready = st.session_state.get('knowledge_base_ready', False)
        
        if not knowledge_base_ready:
            st.info("Please upload and process documents first to enable querying.")
        else:
            st.success("Knowledge base is ready! Ask your questions below.")
        
        # Query input
        user_query = st.text_input(
            "Enter your question:",
            placeholder="How to create a pipeline object?",
            disabled=not knowledge_base_ready
        )
        
        # Query button
        if st.button("Ask Question", disabled=not knowledge_base_ready or not user_query):
            if user_query and rag_system.vector_database:
                with st.spinner("Searching and generating answer..."):
                    answer, retrieved_docs = rag_system.answer_with_rag(
                        user_query, 
                        num_retrieved_docs=num_docs
                    )
                    
                    # Display answer
                    st.subheader("Answer")
                    st.write(answer)
                    
                    # Display retrieved documents
                    st.subheader("Retrieved Documents")
                    for i, doc in enumerate(retrieved_docs):
                        with st.expander(f"Document {i+1}"):
                            st.text_area(
                                "Content", 
                                doc, 
                                height=150, 
                                disabled=True,
                                key=f"doc_{i}"
                            )
        
        # Query history
        if 'query_history' not in st.session_state:
            st.session_state.query_history = []
        
        if user_query and st.button("Save Query", disabled=not user_query):
            st.session_state.query_history.append(user_query)
        
        if st.session_state.query_history:
            st.subheader("Query History")
            for i, query in enumerate(reversed(st.session_state.query_history[-5:])):
                st.text(f"{len(st.session_state.query_history)-i}. {query}")
    
    # Footer with instructions
    st.markdown("---")
    st.markdown("""
    ## How to Use:
    1. **Initialize Models**: Click 'Initialize Models' in the sidebar
    2. **Upload Documents**: Use the file uploader to add PDF, DOCX, or TXT files
    3. **Process Documents**: Click 'Process Documents' to chunk and embed your files
    4. **Ask Questions**: Enter questions in the query interface
    
    ## Free Deployment Options:
    - **Streamlit Cloud**: Connect your GitHub repo with this app
    - **Hugging Face Spaces**: Upload as a Streamlit Space
    - **Railway/Render**: Deploy with minimal configuration
    """)

if __name__ == "__main__":
    # Install required packages notification - UPDATED LIST
    st.markdown("""
    ### Required Dependencies:
    ```bash
    # Core Packages
    pip install streamlit pandas numpy torch transformers accelerate bitsandbytes
    
    # LangChain Ecosystem (Updated modular imports)
    pip install langchain-core langchain-text-splitters langchain-community langchain-huggingface
    
    # Vector Store & Embeddings
    pip install faiss-cpu sentence-transformers 
    
    # Document Loaders
    pip install PyMuPDF python-docx
    ```
    """)
    
    main()